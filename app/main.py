from flask import Flask, render_template, request, send_file, abort
import os
import tempfile
import subprocess
from docx import Document
from pathlib import Path

app = Flask(__name__)

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        text = request.form.get("editor", "")
        fmt = request.form.get("format", "txt")  # txt | docx | doc

        if not text:
            abort(400, "Пустой текст")

        # Создаём временный каталог
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)

            if fmt == "txt":
                file_path = tmpdir / "output.txt"
                file_path.write_text(text, encoding="utf-8")

            elif fmt == "docx":
                file_path = tmpdir / "output.docx"
                doc = Document()
                doc.add_paragraph(text)
                doc.save(file_path)

            elif fmt == "doc":
                # Сначала создаём .docx, а потом конвертируем в .doc через LibreOffice
                docx_path = tmpdir / "output.docx"
                doc = Document()
                doc.add_paragraph(text)
                doc.save(docx_path)

                # Конвертация
                cmd = [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "doc",
                    "--outdir",
                    str(tmpdir),
                    str(docx_path),
                ]
                subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                file_path = tmpdir / "output.doc"

            else:
                abort(400, f"Неизвестный формат {fmt}")

            # Отдаём файл
            return send_file(
                file_path,
                as_attachment=True,
                download_name=file_path.name,
                mimetype="application/octet-stream",
            )

    return render_template("index.html")


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)